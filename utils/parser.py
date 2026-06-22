"""
📄 PARSER — TXT/PDF/DOCX fayllardan savollar ajratish

✅ STANDART FORMAT A:
  1. Savol matni?
  ===A) To'g'ri javob
  B) Xato
  C) Xato
  D) Xato

QO'LLAB-QUVVATLANADIGAN FORMATLAR:
  FORMAT B: ==== + # + ++++ (ZIP fayllar)
  FORMAT C: ? = formati (PDF)
  FORMAT D: Jadval (Ko'p ustunli DOCX)
  FORMAT E: Ha/Yo'q, Bo'sh joy to'ldirish, Erkin javob
  FORMAT J: Markersiz, ketma-ket
"""
import re, logging, os, subprocess, tempfile
from pathlib import Path

log = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════
#  ASOSIY KIRISH NUQTASI
# ═══════════════════════════════════════════════════════════

def check_images_in_file(path: str) -> dict:
    """
    Fayl ichida rasm bor-yo'qligini tekshiradi.
    Qaytaradi: {"has_images": bool, "count": int, "type": "docx/pdf/none"}
    """
    import os
    ext = os.path.splitext(path)[1].lower()
    result = {"has_images": False, "count": 0, "type": "none"}

    try:
        if ext == ".docx":
            from docx import Document
            try:
                doc = Document(path)
                cnt = len(doc.inline_shapes)
                if cnt == 0:
                    import zipfile
                    with zipfile.ZipFile(path) as z:
                        cnt = len([n for n in z.namelist()
                                   if 'media/image' in n and not n.endswith('/')])
                result = {"has_images": cnt > 0, "count": cnt, "type": "docx"}
            except Exception:
                import zipfile
                with zipfile.ZipFile(path) as z:
                    cnt = len([n for n in z.namelist()
                               if 'media/image' in n and not n.endswith('/')])
                result = {"has_images": cnt > 0, "count": cnt, "type": "docx"}
        elif ext == ".pdf":
            try:
                import fitz
                doc = fitz.open(path)
                cnt = sum(len(page.get_images()) for page in doc)
                doc.close()
                result = {"has_images": cnt > 0, "count": cnt, "type": "pdf"}
            except ImportError:
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    cnt = sum(len(p.images) for p in pdf.pages)
                result = {"has_images": cnt > 0, "count": cnt, "type": "pdf"}
    except Exception as e:
        log.warning(f"check_images_in_file: {e}")
    return result


def parse_file(path: str) -> list:
    """✅ FIXED: Standart format A dastur bo'yicha parse qilish"""
    ext = Path(path).suffix.lower()
    try:
        if ext == ".doc":
            path = _convert_doc(path)
            if not path:
                return []
            ext = ".docx"
        if ext == ".docx":
            return _parse_docx(path)
        elif ext == ".pdf":
            return _parse_pdf(path)
        elif ext == ".txt":
            return _parse_txt(path)
        elif ext in (".xlsx", ".xls", ".xlsm"):
            return _parse_xlsx(path)
        else:
            return []
    except Exception as e:
        log.error(f"parse_file xato ({ext}): {e}", exc_info=True)
        return []


def _convert_doc(path: str) -> str:
    """✅ .doc → .docx konvertatsiya"""
    outdir = tempfile.mkdtemp()
    try:
        r = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx",
             path, "--outdir", outdir],
            capture_output=True, timeout=60
        )
        new = os.path.join(outdir, Path(path).stem + ".docx")
        if os.path.exists(new) and os.path.getsize(new) > 100:
            log.info(f"DOC → DOCX (LibreOffice): {Path(path).name}")
            return new
    except Exception as e:
        log.warning(f"LibreOffice yo'q yoki xato: {e}")
    try:
        import docx as _docx
        _docx.Document(path)
        log.info(f"DOC to'g'ridan DOCX sifatida o'qildi: {Path(path).name}")
        return path
    except Exception:
        pass
    try:
        import docx2txt
        txt = docx2txt.process(path)
        if txt and len(txt) > 50:
            txt_path = os.path.join(outdir, Path(path).stem + ".txt")
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(txt)
            log.info(f"DOC → TXT (docx2txt): {Path(path).name}")
            return txt_path
    except Exception:
        pass
    log.error(f"DOC convert qilinmadi: {path}")
    return ""


# ═══════════════════════════════════════════════════════════
#  DOCX PARSER
# ═══════════════════════════════════════════════════════════

def _parse_docx(path: str) -> list:
    try:
        from docx import Document
        doc = Document(path)
    except Exception as e:
        if 'NULL' in str(e):
            log.warning(f"DOCX NULL rel, ZIP fallback: {e}")
            return _parse_docx_via_zip(path)
        log.error(f"DOCX ochilmadi: {e}")
        return []
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if _is_eq_format(lines):
        q = _parse_eq_hash(lines)
        if q:
            return q
    full_text = "\n".join(lines)
    q = parse_text(full_text)
    if q:
        return q
    return _parse_loose_paragraphs(lines)


# ═══════════════════════════════════════════════════════════
#  FORMAT A — STANDART RAQAMLI ✅ FIXED
# ═══════════════════════════════════════════════════════════

def parse_text(text: str) -> list:
    """✅ FIXED: Standart format A
    1. Savol matni?
    ===A) To'g'ri javob
    B) Xato
    C) Xato
    """
    text = text.replace("\r\n", "\n")
    # Savolni 1. 2. 3. raqamlar yoki ===A) bilan topamiz
    blocks = re.split(r"\n(?=\d+[\.)\]]|===)", "\n" + text.strip())
    result = []
    for b in blocks:
        q = _parse_block(b.strip())
        if q:
            result.append(q)
    return result


def _parse_block(block: str) -> dict:
    """✅ Standart FORMAT A: 1. Savol === A) To'g'ri B) Xato..."""
    if not block:
        return None
    lines = [l.rstrip() for l in block.split("\n") if l.strip()]
    if not lines:
        return None
    q_text = re.sub(r"^\d+[\.)\]]\s*", "", lines[0]).strip()
    if not q_text:
        return None
    opts = []
    corr = None
    for line in lines[1:]:
        ls = line.strip()
        if not ls:
            continue
        # ✅ ===A) format
        if ls.startswith("==="):
            rest = ls[3:].strip()
            opts.append(rest)
            corr = rest
        # ✅ A) format
        elif re.match(r"^[A-Ha-h]\s*[\)\.]\s*", ls):
            opts.append(ls)
        else:
            continue
    if len(opts) < 2:
        return None
    return {
        "type": "multiple_choice",
        "question": q_text,
        "options": opts,
        "correct": corr or "",
        "explanation": "",
        "accepted_answers": [],
        "points": 1,
        "_marked": bool(corr),
    }


# ═══════════════════════════════════════════════════════════
#  FORMAT B: ==== + # + ++++ parser
# ═══════════════════════════════════════════════════════════

def _is_eq_format(lines: list) -> bool:
    has_eq   = any(re.match(r"^={3,}$", l) for l in lines)
    has_plus = any(re.match(r"^\+{3,}$", l) for l in lines)
    return has_eq and has_plus


def _parse_eq_hash(lines: list) -> list:
    """FORMAT B: ++++ savol ===  # to'g'ri ====  Xato"""
    LBL = ["A", "B", "C", "D", "E", "F", "G", "H"]
    questions = []
    content = "\n".join(lines)
    blocks = re.split(r'\+{4,}', content)
    blocks = [b.strip() for b in blocks if b.strip()]
    for block in blocks:
        parts = re.split(r'={3,}', block)
        parts = [p.strip() for p in parts if p.strip()]
        if len(parts) < 2:
            continue
        question = parts[0]
        if not question:
            continue
        correct_idx = -1
        clean_answers = []
        for ans in parts[1:]:
            ans = ans.strip()
            if ans.startswith('#'):
                if correct_idx == -1:
                    correct_idx = len(clean_answers)
                clean_answers.append(ans[1:].strip())
            else:
                clean_answers.append(ans)
        clean_answers = [a for a in clean_answers if a]
        if not clean_answers:
            continue
        opts = []
        for i, ans in enumerate(clean_answers):
            lbl = LBL[i] if i < len(LBL) else str(i + 1)
            opts.append(ans if re.match(r"^[A-Ha-h]\s*[\)\.]" , ans) else f"{lbl}) {ans}")
        questions.append({
            "type": "multiple_choice",
            "question": question,
            "options": opts,
            "correct": opts[correct_idx] if correct_idx >= 0 else "",
            "explanation": "",
            "accepted_answers": [],
            "points": 1,
            "_marked": correct_idx >= 0,
        })
    return questions


# ═══════════════════════════════════════════════════════════
#  QOLGAN PARSER'LAR (QISQARTIB KO'RSATILDI)
# ═══════════════════════════════════════════════════════════

def _parse_txt(path: str) -> list:
    try:
        with open(path, encoding="utf-8") as f:
            text = f.read()
    except Exception:
        return []
    return parse_text(text)


def _parse_pdf(path: str) -> list:
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception as e:
        log.error(f"PDF xato: {e}")
        return []
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if _is_eq_format(lines):
        return _parse_eq_hash(lines)
    return parse_text(text)


def _parse_xlsx(path: str) -> list:
    try:
        import pandas as pd
        df = pd.read_excel(path, header=None)
        lines = [str(v).strip() for v in df.iloc[:, 0] if str(v) != "nan"]
    except Exception:
        return []
    return parse_text("\n".join(lines))


def _parse_docx_via_zip(path: str) -> list:
    """✅ NULL relationship xatosi bo'lganda ZIP orqali o'qiydi"""
    try:
        import zipfile
        with zipfile.ZipFile(path) as z:
            doc_xml = z.read('word/document.xml')
        from xml.etree import ElementTree as ET
        root = ET.fromstring(doc_xml)
        text = ""
        for t in root.iter():
            if t.text:
                text += t.text + " "
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        return parse_text(text)
    except Exception as e:
        log.error(f"ZIP parse: {e}")
        return []


def _parse_loose_paragraphs(lines: list) -> list:
    """FORMAT I: Alohida paragraflar"""
    LBL = ["A", "B", "C", "D", "E", "F", "G", "H"]
    if not lines:
        return []
    questions = []
    i = 0
    while i < len(lines):
        if not lines[i] or lines[i].startswith(("=", "+", "#")):
            i += 1
            continue
        q_text = lines[i]
        i += 1
        opts = []
        correct_idx = -1
        while i < len(lines):
            l = lines[i].strip()
            if not l:
                i += 1
                continue
            if l.startswith(("=", "+", "#")) and len(opts) > 0:
                break
            if l.startswith("+"):
                correct_idx = len(opts)
                opts.append(l[1:].strip())
            else:
                opts.append(l)
            i += 1
        if len(opts) >= 2:
            formatted_opts = [f"{LBL[j]}) {o}" if not re.match(r"^[A-Ha-h]\s*[\)\.]" , o) else o for j, o in enumerate(opts[:8])]
            questions.append({
                "type": "multiple_choice",
                "question": q_text,
                "options": formatted_opts,
                "correct": formatted_opts[correct_idx] if correct_idx >= 0 else "",
                "explanation": "",
                "accepted_answers": [],
                "points": 1,
                "_marked": correct_idx >= 0,
            })
    return questions
